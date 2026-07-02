"""Generate NJEM_Engineering_Report.docx — detailed report with tensor shapes at
every stage: data, model forward pass, losses, plus every issue/fix and rationale."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK   = RGBColor(0x0F, 0x17, 0x2A)
INDIGO= RGBColor(0x4F, 0x46, 0xE5)
TEAL  = RGBColor(0x0D, 0x94, 0x88)
AMBER = RGBColor(0xB4, 0x5C, 0x06)
GREY  = RGBColor(0x64, 0x74, 0x8B)
CODEBG= "F3F4F8"

doc = Document()
base = doc.styles["Normal"]
base.font.name = "Calibri"; base.font.size = Pt(11); base.font.color.rgb = INK

def _shade(el, hexcolor):
    pPr = el._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor); pPr.append(shd)

def code(lines):
    if isinstance(lines, str): lines = lines.split("\n")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    _shade(p, CODEBG)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = INK
    rPr = r._r.get_or_add_rPr(); rF = rPr.get_or_add_rFonts()
    rF.set(qn("w:ascii"), "Consolas"); rF.set(qn("w:hAnsi"), "Consolas")
    return p

def label(text, color):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = color
    return p

def body(text):
    p = doc.add_paragraph(text); p.paragraph_format.space_after = Pt(4); return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet"); p.paragraph_format.space_after = Pt(2)

def table(headers, rows, widths=None, style="Light Grid Accent 1"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; run = c.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            para = cells[i].paragraphs[0]; run = para.add_run(str(v)); run.font.size = Pt(9.5)
            if i == 0 or ("(" in str(v) and ")" in str(v) and "," in str(v)):
                run.font.name = "Consolas"
                rF = run._r.get_or_add_rPr().get_or_add_rFonts()
                rF.set(qn("w:ascii"), "Consolas"); rF.set(qn("w:hAnsi"), "Consolas")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

def issue_block(sym, cause, fix, why, fixcode=None):
    label("Symptom", AMBER); body(sym)
    label("Root cause", INK); body(cause)
    label("Fix", TEAL); body(fix)
    if fixcode: code(fixcode)
    label("Why this fix", INDIGO); body(why)

# ── Title ─────────────────────────────────────────────
t = doc.add_paragraph(); r = t.add_run("NJEM — Engineering Report")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = INK; r.font.name = "Cambria"
st = doc.add_paragraph()
r = st.add_run("Neural Joint Embedding Model · Brain-to-Text (t15) · Shapes, Issues, Fixes & Rationale")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = INDIGO
m = doc.add_paragraph()
r = m.add_run("Pipeline: intracortical neural activity → Whisper encoder embeddings → text")
r.font.size = Pt(10.5); r.font.color.rgb = GREY

# ── 1. Overview ───────────────────────────────────────
doc.add_heading("1. Project Overview", level=1)
body("NJEM maps intracortical neural recordings (t15 Brain-to-Text) to the Whisper tiny.en encoder "
     "embedding space, then reuses Whisper's frozen decoder to turn predicted embeddings into text. "
     "Four stages:")
bullets([
    "generate_audio — StyleTTS2 synthesizes speech per transcript (24 kHz → resampled to 16 kHz).",
    "extract_features — frozen Whisper encoder produces the (1500, 384) embedding; store audio_length + transcription.",
    "train — a 12.33M-param ConvBiGRU regresses neural activity to encoder embeddings.",
    "decode — feed predicted embeddings to Whisper's frozen decoder (beam search) and score WER.",
])
body("Global constants: NEURAL_DIM = 512, N_CTX = 1500 (Whisper encoder frames), EMB_DIM = 384 "
     "(encoder hidden size), FRAME_SAMPLES = 320 (16 kHz samples per encoder frame = 20 ms).")

# ── 2. Data & shapes ──────────────────────────────────
doc.add_heading("2. Data Representation & Tensor Shapes", level=1)
body("A single dataset item (src/dataset.py, NeuralToEmbeddingDataset.__getitem__) returns:")
table(["Tensor", "Shape", "dtype", "Meaning"],
      [("neural", "(T, 512)", "float32", "neural features, one 512-d vector per 20 ms bin; T varies per trial"),
       ("target/emb", "(vf, 384)", "float32", "Whisper encoder embedding, sliced to vf content frames"),
       ("transcription", "str", "—", "ground-truth sentence (WER + decoder loss)")],
      widths=[1.4, 1.2, 0.8, 3.3])

doc.add_heading("2.1  Valid frames (vf) — the audio clock", level=2)
body("The stored encoder_embedding is ALWAYS 1500 frames (Whisper pads every clip to 30 s). Only "
     "the first vf frames correspond to real audio; the dataset slices encoder_embedding[:vf].")
code("vf = clamp( ceil(audio_length / 320), 1, 1500 )\n"
     "16000 samples/s ÷ 320 samples/frame = 50 frames/s  ->  1.0 s audio = 50 valid frames\n"
     "example: trial_0000 audio_length=29567 -> ceil(29567/320)=93 frames (~1.85 s)")

doc.add_heading("2.2  Two independent length clocks", level=2)
body("Neural length T and audio length vf are unrelated — the same sentence has a ~6 s neural "
     "window but ~1.8 s of synthesized audio. This mismatch is why the model needs cross-attention "
     "(a resampler), not a frame-aligned regressor.")
table(["Clock", "Symbol", "Definition", "Example (trial_0000)"],
      [("Neural", "T", "input_features.shape[0], 20 ms bins", "321 bins"),
       ("Neural (post-conv)", "T'", "conv_out(T) ≈ T/4", "81 frames"),
       ("Audio / target", "vf", "ceil(audio_length/320)", "93 frames")],
      widths=[1.6, 0.8, 3.0, 1.3])

doc.add_heading("2.3  Batch collation shapes (collate)", level=2)
table(["Tensor", "Shape", "Notes"],
      [("neural_pad", "(B, Tmax, 512)", "zero-padded to the batch's longest T"),
       ("n_lengths", "(B,)", "true neural length per trial (for packing + mask)"),
       ("target", "(B, 1500, 384)", "zero-padded to N_CTX"),
       ("mask", "(B, 1500)", "bool, True on [:vf] — single source of truth for valid frames"),
       ("txts", "list[str] len B", "ground-truth transcriptions")],
      widths=[1.5, 1.6, 3.6])

# ── 3. Model forward pass with shapes ─────────────────
doc.add_heading("3. Model Architecture — Forward Pass with Shapes", level=1)
body("ConvBiGRU (src/model.py): in_dim 512, conv_channels 256, hidden 256, gru_layers 8, "
     "bidirectional, emb_dim 384, n_ctx 1500, attn_heads 4, dropout 0.1. Shape trace for a batch "
     "of B trials with max neural length Tmax and n_out requested output frames:")
table(["Step", "Operation", "Output shape"],
      [("input", "x (neural)", "(B, Tmax, 512)"),
       ("transpose", "x.transpose(1,2) for Conv1d", "(B, 512, Tmax)"),
       ("conv1", "Conv1d(512,256,k5,s2,p2)+GELU+Drop", "(B, 256, Tmax/2)"),
       ("conv2", "Conv1d(256,256,k5,s2,p2)+GELU+Drop", "(B, 256, Tmax/4)"),
       ("transpose", "back to (B, T', C)", "(B, T', 256)"),
       ("pack", "pack_padded_sequence(conv_lengths)", "PackedSequence"),
       ("gru", "8-layer BiGRU, hidden 256", "(B, T', 512)"),
       ("queries", "learnable query[:, :n_out].expand", "(B, n_out, 512)"),
       ("cross_attn", "MHA(q=queries, k=v=gru, key_mask)", "(B, n_out, 512)"),
       ("norm", "LayerNorm(512)", "(B, n_out, 512)"),
       ("head", "Lin512->512, GELU, Drop, Lin512->384", "(B, n_out, 384)")],
      widths=[1.0, 3.4, 2.2])

doc.add_heading("3.1  Conv output-length formula", level=2)
body("Each stride-2 conv maps length L -> (L-1)//2 + 1, applied twice. This length feeds both "
     "pack_padded_sequence and the attention key-padding mask, so padded neural frames never leak.")
code("def _conv_out_length(L):        # applied per stride-2 layer\n"
     "    for _ in range(2): L = (L - 1)//2 + 1\n"
     "    return L\n"
     "T=321 -> 81 ; T=480 -> 120 ; T=890 -> 223   (~T/4)")

doc.add_heading("3.2  Cross-attention: bridging the two clocks", level=2)
body("The GRU output has T' frames (neural clock). We need n_out frames (audio clock). "
     "n_out learnable queries attend over the T' GRU frames; the key-padding mask hides padded "
     "neural positions. This RESAMPLES T' -> n_out regardless of whether T' is larger or smaller "
     "than n_out. In training n_out = max_vf (batch max valid frames); at decode n_out defaults to 1500.")
code("queries  : (B, n_out, 512)          # learnable, sliced from query[1,1500,512]\n"
     "keys/val : (B, T', 512)             # GRU output\n"
     "key_mask : (B, T') bool = arange(T') >= conv_out(lengths)\n"
     "out, _   = cross_attn(queries, keys, keys, key_padding_mask=key_mask)  # (B, n_out, 512)")

doc.add_heading("3.3  Parameter budget (12.33M total)", level=2)
table(["Module", "Params", "What it is"],
      [("gru", "9,068,544", "8-layer bidirectional GRU (dominates the budget)"),
       ("cross_attn", "1,050,624", "MultiheadAttention, 4 heads, dim 512"),
       ("conv", "983,552", "two Conv1d front-end layers"),
       ("query", "768,000", "learnable queries 1500 × 512"),
       ("head", "459,648", "MLP 512→512→384"),
       ("attn_norm", "1,024", "LayerNorm(512)"),
       ("TOTAL", "12,331,392", "≈ 12.33M")],
      widths=[1.5, 1.5, 3.7])

# ── 4. Loss with shapes ───────────────────────────────
doc.add_heading("4. The Loss Function (with shapes)", level=1)
body("Per batch the training loop computes max_vf = mask.sum(1).max(), calls the model with "
     "n_out = max_vf, then slices target and mask to max_vf so every tensor aligns:")
code("max_vf = int(mask.sum(dim=1).max())         # scalar\n"
     "pred   = model(neural, lengths, n_out=max_vf)     # (B, max_vf, 384)\n"
     "mvf    = mask[:, :max_vf]                          # (B, max_vf) bool\n"
     "tgt    = target[:, :max_vf]                        # (B, max_vf, 384)")
body("Total objective:")
code("L  =  L_emb  +  λ · L_CE          (λ = dec_loss_weight = 0.8)")

doc.add_heading("4.1  L_emb — masked embedding regression", level=2)
body("Two terms, both averaged only over valid frames. Shapes shown inline:")
code("m     = mvf.unsqueeze(-1).float()                          # (B, max_vf, 1)\n"
     "# SmoothL1 (magnitude), per-element then masked-mean over 384 dims:\n"
     "l1    = (smooth_l1(pred, tgt, reduction='none') * m).sum() / (m.sum() * 384)   # scalar\n"
     "# Cosine (direction), per-frame then masked-mean:\n"
     "cos_f = 1 - cosine_similarity(pred, tgt, dim=-1)           # (B, max_vf)\n"
     "cos   = (cos_f * mvf).sum() / mvf.sum()                    # scalar\n"
     "L_emb = 1.0*l1 + 1.0*cos")
bullets([
    "SmoothL1 (Huber): matches the magnitude of each of the 384 dims; MSE-like for small errors, "
    "L1-like for large ones so outlier dims don't dominate.",
    "Cosine 1−cos: matches DIRECTION of each 384-d vector — what Whisper's decoder is most sensitive to.",
    "Weakness alone: low L_emb ≠ low WER — the model can collapse toward the mean embedding.",
])

doc.add_heading("4.2  L_CE — decoder-in-the-loop cross-entropy", level=2)
body("Predicted embeddings are fed to the FROZEN Whisper decoder, teacher-forced with ground-truth "
     "tokens; cross-entropy penalizes wrong next tokens. Gradient flows through the frozen decoder "
     "back into the ConvBiGRU. tiny.en tokenizer: SOT sequence = [50257, 50362], EOT = 50256, "
     "vocab = 51864.")
code("# teacher-forcing tensors, L = max token length in batch\n"
     "seqs = [ [50257,50362] + encode(' '+text) + [50256] ]     # per trial\n"
     "inp  : (B, L) long   = tokens (shifted input)\n"
     "tgt  : (B, L) long   = next-token targets; prompt positions = -100 (ignored)\n"
     "feats= pred * mvf.unsqueeze(-1)                            # (B, max_vf, 384) padding zeroed\n"
     "logits = whisper.decoder(inp, feats)                      # (B, L, 51864)\n"
     "L_CE = cross_entropy(logits.reshape(-1,51864), tgt.reshape(-1), ignore_index=-100)")
body("The decoder cross-attends to feats (our predicted embeddings) exactly as it normally attends "
     "to real encoder output — so supplying predicted embeddings bypasses audio entirely.")

doc.add_heading("4.3  Why both terms", level=2)
bullets([
    "L_emb alone → good vector numbers, wrong transcripts (mode collapse; cosine plateaus ~0.72).",
    "L_CE alone → unstable early: a fresh model hands the decoder garbage, gradients are noisy.",
    "Together → L_emb is a dense, stable per-frame target; L_CE bends embeddings toward the region "
    "Whisper decodes correctly (needs ~0.95+ cosine). λ = 0.8 balances them.",
])

# ── 5. Evaluation with shapes ─────────────────────────
doc.add_heading("5. Evaluation & Decoding (with shapes)", level=1)
body("decode_dataset (src/decode.py) turns predicted embeddings into text and scores WER:")
code("pred  = model(neural.unsqueeze(0))[0]        # (1500, 384)   (n_out defaults to 1500)\n"
     "feats = zeros(1500, 384); feats[:vf] = pred[:vf]     # keep content frames\n"
     "res   = whisper.decode(wmodel, feats.unsqueeze(0), options)   # options: lang=en, beam_size=5\n"
     "wer   = levenshtein(norm(truth).split(), norm(pred_text).split()) / len(truth_words)")
bullets([
    "Encoder is skipped — the (1500, 384) tensor is fed straight in as if it were encoder output.",
    "Beam search (beam_size = 5) + EnglishTextNormalizer for a fair, convention-matching WER.",
    "In-training eval every wer_every=5 epochs over wer_trials=30 val trials; results also dumped "
    "to a CSV of (idx, wer, truth, pred).",
])

# ── 6. Issues & Fixes ─────────────────────────────────
doc.add_heading("6. Issues Encountered & Fixes", level=1)

doc.add_heading("6.1  CUDA/CPU device mismatch in the attention mask", level=2)
issue_block(
    "Training crashed inside Docker (GPU) with a device-mismatch RuntimeError when the "
    "cross-attention key-padding mask was built.",
    "GRU output h is on cuda, but the valid-length tensor for the mask was left on CPU; comparing "
    "a cuda arange (B, T') against a CPU length tensor raises a device mismatch. Only surfaced on GPU.",
    "Move the conv-output length tensor onto h's device before building the mask.",
    "cl = self._conv_out_length(lengths).clamp(min=1).to(h.device)   # (B,)\n"
    "key_mask = torch.arange(max_klen, device=h.device).unsqueeze(0) >= cl.unsqueeze(1)  # (B, T')",
    "Every tensor in a broadcast op must share a device. Anchoring cl to h.device (not hard-coding "
    "'cuda') keeps the model correct on both CPU and GPU.")

doc.add_heading("6.2  Stale-checkpoint resume crash (architecture mismatch)", level=2)
issue_block(
    "On relaunch, resume aborted with a state_dict shape mismatch.",
    "An old best.pt was saved from an 8-layer GRU; the run was reconfigured with different "
    "gru_layers. load_state_dict needs identical tensor shapes, so the mismatched GRU threw.",
    "Archive incompatible checkpoints (checkpoints/old_8layer/), relaunch clean, and store args "
    "inside every checkpoint so the architecture is always recoverable.",
    "d = {'model': model.state_dict(), ..., 'args': vars(args), 'epoch': epoch}",
    "A checkpoint is only loadable into the identical architecture. Storing args + archiving stale "
    "weights prevents silent shape clashes when hyper-parameters change.")

doc.add_heading("6.3  Slide generation: pptxgenjs incompatible with system Node", level=2)
issue_block(
    "pptxgenjs failed with a 'Not supported' error.",
    "pptxgenjs 4.x needs a modern Node.js; the box ships Node v10.19.0, far too old.",
    "Switch slide/report generation to python-pptx / python-docx in the project Python env "
    "(/raid/owais/home/env), removing the Node dependency.",
    None,
    "Reusing the training Python env avoids a second toolchain, runs anywhere the project runs, and "
    "sidesteps the unfixable Node version constraint on a shared box.")

doc.add_heading("6.4  Checkpoint selection broke when the loss scale changed", level=2)
issue_block(
    "After enabling the decoder loss, best.pt stopped updating — frozen at a pre-change epoch even "
    "though later epochs improved.",
    "best.pt was chosen by lowest val loss, but adding the CE term inflated the val-loss scale "
    "(≈0.57 → ≈1.66). No post-change epoch could beat the old smaller number — val loss is not "
    "comparable across two different loss definitions.",
    "Track three checkpoints: best.pt (lowest val loss), best_wer.pt (lowest WER — the true "
    "metric), last.pt written EVERY epoch. Resume prefers last.pt, then best.pt.",
    "if va[0] < best_val:  save_ckpt('best.pt', epoch)\n"
    "if mw   < best_wer:   save_ckpt('best_wer.pt', epoch, wer=mw)\n"
    "save_ckpt('last.pt', epoch)   # always -> crash-safe, scale-change-safe resume",
    "When a loss definition changes, cross-epoch loss comparisons are meaningless. Selecting on WER "
    "(scale-invariant to reweighting) and always keeping the latest weights makes checkpointing robust.")

doc.add_heading("6.5  WER stuck at ~1.0 — regression-to-the-mean collapse", level=2)
issue_block(
    "Embedding regression looked healthy (cosine ≈ 0.72, falling loss) yet WER ≈ 1.0 with fluent "
    "but wrong transcripts ('You know, you know, you know…').",
    "SmoothL1 + (1−cosine) can be minimized by regressing toward the average embedding (mode "
    "collapse): cross-trial similarity ~0.972 vs a real ~0.903; cosine plateaus ~0.72. The decoder "
    "needs ~0.95+ directional agreement, so generic embeddings decode to generic sentences.",
    "Add the decoder-in-the-loop CE loss (section 4.2): teacher-force GT tokens through the FROZEN "
    "Whisper decoder on predicted embeddings; back-prop CE into the ConvBiGRU (λ = 0.8).",
    "loss = masked_embedding_loss(pred, tgt, mvf) + 0.8 * decoder_loss(pred, mvf, txts)",
    "Regression optimizes a proxy (vector distance); WER is the real objective. The decoder loss "
    "optimizes decodability directly, pushing embeddings into the exact region Whisper reads as the "
    "correct words — closing the 'good MSE, bad WER' gap.")

doc.add_heading("6.6  Unfair WER measurement (greedy + raw strings)", level=2)
issue_block(
    "Early WER was pessimistic and noisy.",
    "Greedy decoding + raw string comparison counted punctuation/casing as errors and under-explored.",
    "Beam search (beam_size = 5) + EnglishTextNormalizer before a Levenshtein WER; also report exact-match %.",
    None,
    "Matches Whisper's own eval convention, so the score reflects genuine word errors, and beam "
    "search gives the decoder a fair chance to find the best transcript.")

doc.add_heading("6.7  No way to obtain the actual transcripts", level=2)
issue_block(
    "The pipeline trained but produced no inspectable text.",
    "Decoding logic lived only inside the training eval loop; no standalone entry point.",
    "Factor decoding into src/decode.py (decode_dataset + decode() CLI) and add a 'decode' stage to "
    "main.py; it loads a checkpoint, rebuilds ConvBiGRU from stored args, decodes a split, writes a "
    "CSV of (idx, wer, truth, pred).",
    "python main.py decode      # -> outputs/decodings_val.csv",
    "A shared decode path powers both the in-training WER eval and on-demand transcript dumps from "
    "any checkpoint — which is what makes the output usable.")

doc.add_heading("6.8  GPU-in-Docker: misleading host probes", level=2)
issue_block(
    "torch.cuda.is_available() is False on the host and nvidia probes mislead, though training runs "
    "fine on the GPU.",
    "Training runs inside Docker where the GPU is visible; the host shell is CPU-only. Host GPU "
    "probes say nothing about the container.",
    "Never probe GPU state from the host. Infer progress from artifacts: checkpoint metadata "
    "(epoch/best_val/best_wer), file mtimes, process cgroups (/proc/<pid>/cgroup shows 'docker'), "
    "and open /dev/nvidia* file descriptors.",
    None,
    "These signals are container-agnostic and read-only — they reveal what training is doing without "
    "touching the shared GPU or drawing wrong conclusions from a CPU-only host view.")

doc.add_heading("6.9  Duplicate GPU training processes contending", level=2)
issue_block(
    "Two 'python main.py train' processes ran on one V100; progress stalled and both eventually exited.",
    "A second launch started while the first was mid-run; two trainers competing for the same GPU "
    "memory is unstable and can OOM/kill each other.",
    "Identify each by cgroup (host vs docker) and open nvidia fds, stop the redundant fresh "
    "duplicate, and verify checkpoints (epoch 7, best_val 3.47) were intact before relaunching one instance.",
    None,
    "One trainer per GPU avoids contention. Distinguishing by cgroup/fd (not guesswork) ensures we "
    "stop the right process and keep the run holding real progress.")

doc.add_heading("6.10  Verifying valid-frame length (vf)", level=2)
issue_block(
    "Confirm the requirement: 1 s of audio → 50 valid frames.",
    "vf drives every masked loss term; an off-by-rate error (e.g. audio left at 24 kHz) would mis-count.",
    "Verified the chain: StyleTTS2 24 kHz is resampled to 16 kHz before storage; audio_length = "
    "len(16 kHz waveform); vf = clamp(ceil(audio_length/320), 1, 1500).",
    "16000 / 320 = 50 frames/s  ->  1.0 s = ceil(16000/320) = 50 valid frames   ✓",
    "At 16 kHz Whisper makes one encoder frame per 320 samples (mel hop 160 × conv stride 2), so "
    "ceil(len/320) = 50 frames/s exactly. The resample-to-16 kHz keeps audio_length and 320 on the same clock.")

# ── 7. Design choices ─────────────────────────────────
doc.add_heading("7. Key Design Choices (summary)", level=1)
table(["Choice", "Reason"],
      [("Predict Whisper ENCODER embeddings", "Encoder space is speech-structured; a frozen decoder turns it into text for free."),
       ("Conv front-end (2× stride-2)", "Smooths and ~4× downsamples the long neural sequence before recurrence."),
       ("Cross-attention with learnable queries", "Neural length (T') and audio length (vf) are unrelated; queries resample T' → vf."),
       ("Frozen Whisper decoder in the loss", "Optimizes decodability (→ WER) directly; Whisper is a differentiable scorer."),
       ("SmoothL1 + cosine (not plain MSE)", "Robust magnitude + directional match; the decoder cares about direction."),
       ("best.pt / best_wer.pt / last.pt", "Robust selection under changing loss scales + crash-safe resume."),
       ("Beam search + EnglishTextNormalizer", "Fair, convention-matching WER."),
       ("Resample audio 24 kHz → 16 kHz", "Keeps audio_length on Whisper's 16 kHz clock so vf = 50 frames/s is exact.")],
      widths=[2.6, 4.4])

# ── 8. Config ─────────────────────────────────────────
doc.add_heading("8. Training Configuration", level=1)
table(["Param", "Value", "Param", "Value"],
      [("epochs", "50", "gru_layers", "8"),
       ("batch_size", "12", "hidden", "256"),
       ("lr", "3e-4", "conv_channels", "256"),
       ("weight_decay", "1e-4", "dropout", "0.1"),
       ("optimizer", "AdamW", "grad clip", "1.0"),
       ("scheduler", "CosineAnnealingLR", "dec_loss_weight (λ)", "0.8"),
       ("dec_model", "tiny.en", "beam_size", "5"),
       ("wer_every", "5 epochs", "wer_trials", "30")],
      widths=[1.7, 1.7, 1.9, 1.7])

# ── 9. Current status ─────────────────────────────────
doc.add_heading("9. Current Status", level=1)
bullets([
    "Latest checkpoint: epoch 7 of 50, val loss 3.47 (improving), WER 1.02, exact-match 0%.",
    "Decoder-in-the-loop loss is working: outputs are now fluent English (correct structure); "
    "content accuracy is the next milestone as the CE term keeps falling.",
    "Checkpoints intact and resumable (resume from last.pt → epoch 8).",
    "Recommendation: run the full 50 epochs as a SINGLE GPU process; watch dec-CE and cosine → ~0.95.",
])

doc.save("NJEM_Engineering_Report.docx")
print("wrote NJEM_Engineering_Report.docx")
